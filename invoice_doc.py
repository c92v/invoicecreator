from docx import Document
from docx.shared import Inches
from datetime import date
import json
import os.path

class InvoiceDocument:
        def __init__(self):
            self.document = Document()
            self.invoice_number = ''
            self.date = ''
        def format_header(self):
            self.document.add_heading('''INSERT NAME\nINSERT HOME ADDRESS\nCITY, STATE ZIPCODE\n''')
            p = self.document.add_paragraph('')
            p.add_run('INVOICE\n').bold = True
            p.add_run(self.date)
            p.add_run('''\nINSERT COMPANY\nINSERT STREET ADDRESS HERE\nCITY, STATE ZIPCODE\n''')
            self.document.add_heading('BALANCE DUE (USD): $' + 'ADD_BALANCE_HERE', level=1)
        # TODO: Probably a better way to do to this...
        # ... also, need to increment day by one before setting it to self
        def get_date(self):
            today = date.today().isoformat()
            today = today.split('-')
            today = today[2] + '-' + today[1] + '-' + today[0]
            self.date = today
        # TODO: Increment number by one in .json file
        def get_invoice_number(self):
            with open("number.json", 'r') as f:
                variables = json.load(f)
            self.invoice_number = str(variables["invoice_number"])
        def add_files(self, files):
            table = self.document.add_table(rows=1, cols=5) 
            col_cells = table.rows[0].cells
            col_cells[0].text = "ITEM"
            col_cells[1].text = "DESCRIPTION"
            col_cells[2].text = "UNIT COST"
            col_cells[3].text = "QUANTITY"
            col_cells[4].text = "LINE TOTAL"
            for file_name, values in files.items():
                row_cells = table.add_row().cells
                row_cells[0].text = "Transcription"
                # TODO: Format names better
                row_cells[1].text = file_name
                row_cells[2].text = "0.87"
                row_cells[3].text = str(values[0])
                row_cells[4].text = str(values[1])
        def add_footer(self):
            p = self.document.add_paragraph('')
            p.add_run("Total\t$" + "INSERT_TOTAL_HERE\n").bold = True
            p.add_run("Balance Paid\t$0.00\n")
            p.add_run("Balance Due (USD)\t$)" + "INSERT_TOTAL_HERE\n").bold = True
        def save(self):
            save_path = "C:\\Users\\User\\Downloads\\INSERT-NAME_Invoice_" + self.date + ".docx"
            if not os.path.isfile(save_path):
                self.document.save(save_path)
                print("Invoice .docx has been created and saved.")
            else:
                print("Error: File for this week already exists!")
        def create_document(self, files):
            self.get_date()
            self.get_invoice_number()
            self.format_header()
            self.add_files(files)
            self.add_footer()
            self.save()
